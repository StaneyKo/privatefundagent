"""按指定私募投顾推荐报告模板生成可直接交付客户的 Word。"""

from __future__ import annotations

import math
import re
import zipfile
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image

from ai.research_grouping import COMMON_GROUP, split_grouped_text, split_strategy_names
from config.config import CONTENT_TEMPLATE_PATH, OUTPUT_DIR
from document.txt_generator import _safe_filename
from models import GeneratedResearch, ResearchClaim, ResearchFact, ResearchSection, SourceChunk


FONT_NAME = "宋体"
TEXT = "000000"
MUTED = "666666"
ACCENT = "17365D"
TABLE_TOTAL_DXA = 8304
SECTION_LABELS = {
    "core_overview": "公司与产品概览",
    "strategy_framework": "策略框架与收益来源",
    "risk_control": "风险控制与组合约束",
    "performance": "代表产品表现",
    "evaluation": "综合评价",
    "market_environment": "适用与不利市场环境",
    "other_information": "其他信息",
}
CLIENT_HIDDEN_SECTIONS = {"internal_business", "follow_up", "other_information"}
CHAPTER_NUMERALS = ("一", "二", "三", "四", "五", "六", "七", "八", "九", "十")


def generate_word(
    company: str,
    strategy: str,
    research: GeneratedResearch | str,
    output_dir: Path = OUTPUT_DIR,
    source_chunks: Iterable[SourceChunk] | None = None,
    performance_chart: dict | None = None,
) -> tuple[bytes, Path]:
    """使用模板纸张与宋体段落格式生成无引用标记的客户版 DOCX。"""
    del source_chunks  # 来源仅供后台核验，不进入客户文件。
    output_dir.mkdir(parents=True, exist_ok=True)
    structured = research if isinstance(research, GeneratedResearch) else _legacy_research(str(research))
    document = _load_content_template()
    _clear_body(document)
    _configure_document(document)
    _add_title(document, company, strategy)
    _add_metadata(document, structured)

    strategies = split_strategy_names(strategy)
    if len(strategies) > 1:
        _add_multi_strategy_sections(document, structured, strategies)
    else:
        _add_single_strategy_sections(document, structured)

    if structured.performance:
        if len(strategies) > 1:
            _add_multi_performance_tables_at_end(document, structured, strategies)
        else:
            _add_performance_table_at_end(document, structured)
    if performance_chart:
        _add_performance_visual_at_end(document, performance_chart, strategy)
    _clear_headers_and_footers(document)
    _set_fixed_core_properties(document, company, strategy)

    raw = BytesIO()
    document.save(raw)
    data = _normalize_docx_package(raw.getvalue())
    path = output_dir / f"{_safe_filename(company)}_{_safe_filename(strategy)}_详细介绍.docx"
    path.write_bytes(data)
    return data, path


def _load_content_template() -> DocumentObject:
    return Document(str(CONTENT_TEMPLATE_PATH)) if CONTENT_TEMPLATE_PATH.exists() else Document()


def _clear_body(document: DocumentObject) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def _configure_document(document: DocumentObject) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)

    normal = document.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    _set_style_fonts(normal._element, FONT_NAME)
    paragraph_format = normal.paragraph_format
    paragraph_format.line_spacing = 1.5
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)


def _add_title(document: DocumentObject, company: str, strategy: str) -> None:
    company_paragraph = document.add_paragraph()
    company_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    company_paragraph.paragraph_format.first_line_indent = Pt(0)
    company_paragraph.paragraph_format.line_spacing = 1.15
    company_paragraph.paragraph_format.space_before = Pt(0)
    company_paragraph.paragraph_format.space_after = Pt(3)
    company_paragraph.paragraph_format.keep_with_next = True
    _format_run(company_paragraph.add_run(company), 18, bold=True, color=ACCENT)

    strategy_paragraph = document.add_paragraph()
    strategy_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    strategy_paragraph.paragraph_format.first_line_indent = Pt(0)
    strategy_paragraph.paragraph_format.line_spacing = 1.15
    strategy_paragraph.paragraph_format.space_before = Pt(0)
    strategy_paragraph.paragraph_format.space_after = Pt(5)
    strategy_paragraph.paragraph_format.keep_with_next = True
    _format_run(strategy_paragraph.add_run(f"{strategy}策略介绍材料"), 20, bold=True, color=ACCENT)


def _add_metadata(document: DocumentObject, research: GeneratedResearch) -> None:
    visible_labels = ("材料日期", "业绩截止")
    fact_map = {
        fact.label: fact
        for fact in research.metadata
        if fact.value and fact.value != "资料未披露"
    }
    facts = [fact_map[label] for label in visible_labels if label in fact_map]
    if not facts:
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(12)
    paragraph.paragraph_format.keep_with_next = True
    for index, fact in enumerate(facts):
        label_run = paragraph.add_run(f"{fact.label}：")
        _format_run(label_run, 9, bold=True, color=MUTED)
        value_run = paragraph.add_run(fact.value)
        _format_run(value_run, 9, color=MUTED)
        separator = "" if index == len(facts) - 1 else "　｜　"
        separator_run = paragraph.add_run(separator)
        _format_run(separator_run, 9, color=MUTED)


def _add_single_strategy_sections(document: DocumentObject, research: GeneratedResearch) -> None:
    """单策略按独立编号章节输出，正文不再使用行内标签。"""
    chapter = 0
    for section in _sections_for_export(research):
        if section.key in CLIENT_HIDDEN_SECTIONS:
            continue
        texts = _section_texts(section, research.performance)
        if not texts:
            continue
        chapter += 1
        _add_section_heading(document, chapter, SECTION_LABELS.get(section.key, _clean_section_title(section.title)))
        for text in texts:
            _add_claim_paragraphs(document, text)


def _sections_for_export(research: GeneratedResearch) -> list[ResearchSection]:
    """详版章节缺失时，复用已通过核验的简版事实补齐核心内容。"""
    sections = [ResearchSection(item.key, item.title, list(item.claims)) for item in research.sections]
    detail_count = sum(
        len(item.claims)
        for item in sections
        if item.key not in {"performance", *CLIENT_HIDDEN_SECTIONS}
    )
    if detail_count >= 6:
        return sections
    by_key = {item.key: item for item in sections}
    existing_text = {claim.text.strip() for item in sections for claim in item.claims}
    for claim in research.brief_claims:
        if claim.text.strip() in existing_text:
            continue
        key = _brief_section_key(claim.text, bool(research.performance))
        if key is None or key not in by_key:
            continue
        limits = {"core_overview": 3, "strategy_framework": 4, "risk_control": 3, "performance": 1, "other_information": 2}
        if len(by_key[key].claims) >= limits.get(key, 1):
            continue
        by_key[key].claims.append(claim)
        existing_text.add(claim.text.strip())
    return sections


def _brief_section_key(text: str, has_performance_table: bool) -> str | None:
    if any(keyword in text for keyword in ("管理费", "业绩报酬", "费率", "收费", "开放", "赎回", "封盘")):
        return "other_information"
    if any(keyword in text for keyword in ("回撤", "夏普", "累计收益", "年化收益", "超额收益", "业绩方面")):
        return None if has_performance_table else "performance"
    if any(keyword in text for keyword in ("选股", "因子", "持仓", "换手", "模型", "信号", "Alpha", "交易标的")):
        return "strategy_framework"
    if any(keyword in text for keyword in ("风险", "风控", "敞口", "止损", "行业偏离", "个股权重")):
        return "risk_control"
    return "core_overview"


def _add_multi_strategy_sections(
    document: DocumentObject,
    research: GeneratedResearch,
    strategies: list[str],
) -> None:
    """先写共同章节，再以每个策略为独立章节，策略名只出现一次。"""
    sections = [
        section
        for section in _sections_for_export(research)
        if section.key not in CLIENT_HIDDEN_SECTIONS
    ]
    chapter = 0
    for section in sections:
        common_claims = _claims_for_group(section.claims, COMMON_GROUP, strategies, include_unmarked=True)
        common_facts = _facts_for_group(research.performance, COMMON_GROUP, strategies, include_unmarked=True)
        texts = _section_texts(ResearchSection(section.key, section.title, common_claims), common_facts)
        if not texts:
            continue
        chapter += 1
        _add_section_heading(document, chapter, SECTION_LABELS.get(section.key, _clean_section_title(section.title)))
        for text in texts:
            _add_claim_paragraphs(document, text)

    for strategy in strategies:
        subsections: list[tuple[str, list[str]]] = []
        for section in sections:
            claims = _claims_for_group(section.claims, strategy, strategies)
            facts = _facts_for_group(research.performance, strategy, strategies)
            texts = _section_texts(ResearchSection(section.key, section.title, claims), facts)
            if not texts:
                continue
            subsections.append((SECTION_LABELS.get(section.key, _clean_section_title(section.title)), texts))
        if not subsections:
            continue
        chapter += 1
        _add_section_heading(document, chapter, strategy)
        for label, texts in subsections:
            _add_subsection_heading(document, label)
            for text in texts:
                _add_claim_paragraphs(document, text)


def _claims_for_group(
    claims: list[ResearchClaim],
    target: str,
    strategies: list[str],
    *,
    include_unmarked: bool = False,
) -> list[ResearchClaim]:
    results: list[ResearchClaim] = []
    for claim in claims:
        group, body = split_grouped_text(claim.text, strategies)
        if group == target or (include_unmarked and group is None):
            results.append(ResearchClaim(body, claim.citations))
    return results


def _facts_for_group(
    facts: list[ResearchFact],
    target: str,
    strategies: list[str],
    *,
    include_unmarked: bool = False,
) -> list[ResearchFact]:
    results: list[ResearchFact] = []
    for fact in facts:
        group, value = split_grouped_text(fact.value, strategies)
        if group == target or (include_unmarked and group is None):
            results.append(ResearchFact(fact.label, value, fact.citations))
    return results


def _add_section_heading(document: DocumentObject, number: int, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.keep_with_next = True
    numeral = CHAPTER_NUMERALS[number - 1] if number <= len(CHAPTER_NUMERALS) else str(number)
    _format_run(paragraph.add_run(f"{numeral}、{title}"), 14, bold=True, color=ACCENT)
    _set_paragraph_bottom_border(paragraph, "9FBAD0", 6, 2)


def _add_subsection_heading(document: DocumentObject, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_before = Pt(7)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.keep_with_next = True
    _format_run(paragraph.add_run(title), 11.5, bold=True, color=ACCENT)


def _section_texts(section: ResearchSection, performance: list[ResearchFact]) -> list[str]:
    texts = [claim.text.strip() for claim in section.claims if claim.text.strip()]
    if section.key != "performance":
        return texts
    summary = _build_performance_summary(performance, " ".join(texts))
    return texts + summary


def _build_performance_summary(facts: list[ResearchFact], existing_text: str) -> list[str]:
    """仅用已核验表格事实补齐产品、数据日期和核心表现描述。"""
    usable = [fact for fact in facts if fact.value and fact.value != "资料未披露"]
    if not usable:
        return []
    existing = existing_text.replace(" ", "")
    product = next(
        (fact for fact in usable if any(keyword in fact.label for keyword in ("代表产品", "产品名称", "基金名称"))),
        None,
    )
    date_fact = next(
        (fact for fact in usable if any(keyword in fact.label for keyword in ("数据时间", "数据日期", "业绩截止", "截止日期", "统计区间"))),
        None,
    )
    overview_parts: list[str] = []
    if product and product.value.replace(" ", "") not in existing:
        overview_parts.append(f"材料披露的代表产品为{product.value}")
    if date_fact and date_fact.value.replace(" ", "") not in existing:
        if "区间" in date_fact.label:
            overview_parts.append(f"业绩统计区间为{date_fact.value}")
        else:
            overview_parts.append(f"{date_fact.label}为{date_fact.value}")

    metric_priority = (
        "运行以来收益",
        "累计收益",
        "年化收益",
        "运行以来超额",
        "累计超额",
        "年化超额",
        "最大回撤",
        "超额最大回撤",
        "超额夏普",
        "夏普",
        "卡玛",
        "波动",
        "单位净值",
        "复权后净值",
        "胜率",
    )
    metrics: list[str] = []
    selected_ids: set[int] = set()
    for keyword in metric_priority:
        for fact in usable:
            if fact is product or fact is date_fact or id(fact) in selected_ids:
                continue
            if keyword not in fact.label or fact.value.replace(" ", "") in existing:
                continue
            metrics.append(f"{fact.label}为{fact.value}")
            selected_ids.add(id(fact))
            break
        if len(metrics) >= 8:
            break

    paragraphs: list[str] = []
    if overview_parts:
        paragraphs.append("，".join(overview_parts) + "。")
    if metrics:
        paragraphs.append("材料披露的主要业绩与风险指标包括：" + "，".join(metrics) + "。")
    return paragraphs


def _add_performance_table_at_end(document: DocumentObject, research: GeneratedResearch) -> None:
    """全部文字结束后放置代表产品指标表，不制造不必要的空白页。"""
    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.line_spacing = 1.5
    heading.paragraph_format.space_before = Pt(16)
    heading.paragraph_format.space_after = Pt(6)
    heading.paragraph_format.keep_with_next = True
    run = heading.add_run("代表产品关键指标")
    _format_run(run, 14, bold=True, color=ACCENT)
    _add_performance_table(document, research.performance)


def _add_multi_performance_tables_at_end(
    document: DocumentObject,
    research: GeneratedResearch,
    strategies: list[str],
) -> None:
    """多策略仍把全部指标表置于正文之后，表内不重复策略名称。"""
    grouped: dict[str, list[ResearchFact]] = {strategy: [] for strategy in strategies}
    common: list[ResearchFact] = []
    for fact in research.performance:
        group, value = split_grouped_text(fact.value, strategies)
        cleaned = ResearchFact(fact.label, value, fact.citations)
        if group in grouped:
            grouped[group].append(cleaned)
        else:
            common.append(cleaned)
    if not common and not any(grouped.values()):
        return

    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.line_spacing = 1.5
    heading.paragraph_format.space_before = Pt(16)
    heading.paragraph_format.space_after = Pt(6)
    heading.paragraph_format.keep_with_next = True
    _format_run(heading.add_run("代表产品关键指标"), 14, bold=True, color=ACCENT)

    table_groups: list[tuple[str, list[ResearchFact]]] = []
    if common:
        table_groups.append(("共同指标", common))
    for strategy in strategies:
        facts = grouped[strategy]
        if not facts:
            continue
        table_groups.append((strategy, facts))
    _add_grouped_performance_table(document, table_groups)


def _add_first_line_paragraph(document: DocumentObject, text: str) -> None:
    paragraph = _new_body_paragraph(document)
    run = paragraph.add_run(text.strip())
    _format_run(run, 10.5)


def _add_hanging_paragraph(document: DocumentObject, text: str) -> None:
    _add_first_line_paragraph(document, text)


def _new_body_paragraph(document: DocumentObject, *, space_after: float = 4) -> object:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.first_line_indent = Pt(21)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.widow_control = True
    paragraph.paragraph_format.keep_together = False
    return paragraph


def _add_claim_paragraphs(document: DocumentObject, text: str, *, label: str | None = None) -> None:
    parts = _split_for_readability(_strip_label_prefix(label, text) if label else text)
    if not parts:
        return
    for index, part in enumerate(parts):
        paragraph = _new_body_paragraph(document)
        if index == 0 and label:
            label_run = paragraph.add_run(f"{label}：")
            _format_run(label_run, 10.5, bold=True)
        run = paragraph.add_run(part)
        _format_run(run, 10.5)


def _split_for_readability(text: str, max_chars: int = 180) -> list[str]:
    value = re.sub(r"\s+", " ", text).strip()
    if len(value) <= max_chars:
        return [value] if value else []
    sentences = [item.strip() for item in re.findall(r"[^。！？]+[。！？]?", value) if item.strip()]
    groups: list[str] = []
    current = ""
    for sentence in sentences:
        if current and len(current) + len(sentence) > max_chars:
            groups.append(current)
            current = sentence
        else:
            current += sentence
    if current:
        groups.append(current)
    return groups


def _strip_label_prefix(label: str | None, text: str) -> str:
    cleaned = text.strip()
    if not label:
        return cleaned
    prefixes = [f"{label}：", f"{label}:"]
    aliases = {"交易风控": ("风控方面：", "风险控制：")}
    prefixes.extend(aliases.get(label, ()))
    for prefix in prefixes:
        if cleaned.startswith(prefix):
            return cleaned[len(prefix) :].lstrip()
    return cleaned


def _merge_section_label(label: str, text: str) -> str:
    cleaned = text.strip()
    if cleaned.startswith((f"{label}：", f"{label}:")):
        return cleaned
    aliases = {"交易风控": ("风控方面：", "风险控制：")}
    for prefix in aliases.get(label, ()):
        if cleaned.startswith(prefix):
            cleaned = cleaned[len(prefix) :].lstrip()
            break
    return f"{label}：{cleaned}"


def _add_performance_table(document: DocumentObject, facts: list[ResearchFact]) -> None:
    row_count = 1 + math.ceil(len(facts) / 2)
    table = document.add_table(rows=row_count, cols=4)
    _set_table_geometry(table, [1704, 2448, 1704, 2448])
    _set_table_borders(table, TEXT, 4)
    _repeat_table_header(table.rows[0])
    for index, header in enumerate(("指标", "数值", "指标", "数值")):
        _fill_table_cell(table.rows[0].cells[index], header, bold=True, fill="D9EAF7")
    for row_index in range(1, row_count):
        for pair_index in range(2):
            fact_index = (row_index - 1) * 2 + pair_index
            label_cell = table.rows[row_index].cells[pair_index * 2]
            value_cell = table.rows[row_index].cells[pair_index * 2 + 1]
            if fact_index < len(facts):
                fact = facts[fact_index]
                _fill_table_cell(label_cell, fact.label, bold=True, fill="F3F6FA")
                _fill_table_cell(value_cell, fact.value)
            else:
                _fill_table_cell(label_cell, "")
                _fill_table_cell(value_cell, "")


def _add_grouped_performance_table(
    document: DocumentObject,
    groups: list[tuple[str, list[ResearchFact]]],
) -> None:
    """在一张表中按策略分段，每个指标值只保留原始内容。"""
    row_count = 1 + sum(1 + math.ceil(len(facts) / 2) for _, facts in groups)
    table = document.add_table(rows=row_count, cols=4)
    _set_table_geometry(table, [1704, 2448, 1704, 2448])
    _set_table_borders(table, TEXT, 4)
    _repeat_table_header(table.rows[0])
    for index, header in enumerate(("指标", "数值", "指标", "数值")):
        _fill_table_cell(table.rows[0].cells[index], header, bold=True, fill="D9EAF7")

    row_index = 1
    for group, facts in groups:
        merged = table.rows[row_index].cells[0].merge(table.rows[row_index].cells[3])
        _fill_table_cell(merged, group, bold=True, fill="E8F1F8")
        row_index += 1
        for start in range(0, len(facts), 2):
            for pair_index in range(2):
                fact_index = start + pair_index
                label_cell = table.rows[row_index].cells[pair_index * 2]
                value_cell = table.rows[row_index].cells[pair_index * 2 + 1]
                if fact_index < len(facts):
                    fact = facts[fact_index]
                    _fill_table_cell(label_cell, fact.label, bold=True, fill="F3F6FA")
                    _fill_table_cell(value_cell, fact.value)
                else:
                    _fill_table_cell(label_cell, "")
                    _fill_table_cell(value_cell, "")
            row_index += 1


def _add_performance_visual_at_end(document: DocumentObject, visual: dict, strategy: str) -> None:
    image_data = visual.get("image_data")
    if not image_data:
        return
    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.line_spacing = 1.5
    heading.paragraph_format.space_before = Pt(10)
    heading.paragraph_format.space_after = Pt(6)
    heading.paragraph_format.keep_with_next = True
    product_name = str(visual.get("product_name") or "").strip()
    title = f"{product_name or strategy}代表产品表现"
    run = heading.add_run(title)
    _format_run(run, 10.5)

    width_inches, height_inches = _fit_image_size(image_data)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()
    run.add_picture(BytesIO(image_data), width=Inches(width_inches), height=Inches(height_inches))


def _fit_image_size(image_data: bytes) -> tuple[float, float]:
    with Image.open(BytesIO(image_data)) as image:
        width, height = image.size
    ratio = width / max(height, 1)
    max_width = 5.72
    max_height = 8.35
    fitted_width = max_width
    fitted_height = fitted_width / max(ratio, 0.01)
    if fitted_height > max_height:
        fitted_height = max_height
        fitted_width = fitted_height * ratio
    return fitted_width, fitted_height


def _fill_table_cell(cell, text: str, *, bold: bool = False, fill: str | None = None) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    _format_run(run, 10.5, bold=bold)
    _set_cell_margins(cell, top=70, start=90, bottom=70, end=90)
    if fill:
        _set_cell_shading(cell, fill)


def _set_cell_shading(cell, fill: str) -> None:
    cell_properties = cell._tc.get_or_add_tcPr()
    shading = cell_properties.first_child_found_in("w:shd")
    if shading is None:
        shading = OxmlElement("w:shd")
        cell_properties.append(shading)
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)


def _set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(TABLE_TOTAL_DXA))
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_column = OxmlElement("w:gridCol")
        grid_column.set(qn("w:w"), str(width))
        grid.append(grid_column)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            tc_width = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_width)
            tc_width.set(qn("w:type"), "dxa")
            tc_width.set(qn("w:w"), str(width))


def _set_table_borders(table, color: str, size: int) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:color"), color)


def _set_cell_margins(cell, **margins: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge in ("top", "start", "bottom", "end"):
        value = margins.get(edge)
        if value is None:
            continue
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _repeat_table_header(row) -> None:
    row_properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    row_properties.append(header)


def _format_run(run, size: float, *, bold: bool = False, color: str = TEXT) -> None:
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attribute}"), FONT_NAME)


def _set_style_fonts(style_element, font_name: str) -> None:
    run_properties = style_element.get_or_add_rPr()
    fonts = run_properties.get_or_add_rFonts()
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attribute}"), font_name)


def _clear_headers_and_footers(document: DocumentObject) -> None:
    for section in document.sections:
        for collection in (section.header, section.footer):
            for paragraph in collection.paragraphs:
                paragraph.clear()
                paragraph_properties = paragraph._p.get_or_add_pPr()
                paragraph_style = paragraph_properties.find(qn("w:pStyle"))
                if paragraph_style is not None:
                    paragraph_properties.remove(paragraph_style)
                border = paragraph_properties.find(qn("w:pBdr"))
                if border is not None:
                    paragraph_properties.remove(border)


def _set_paragraph_bottom_border(paragraph, color: str, size: int, space: int) -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    borders = paragraph_properties.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        paragraph_properties.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)


def _clean_section_title(title: str) -> str:
    return re.sub(r"^[一二三四五六七八九十]+、", "", title).strip()


def _set_fixed_core_properties(document: DocumentObject, company: str, strategy: str) -> None:
    fixed_time = datetime(2000, 1, 1, 0, 0, 0)
    properties = document.core_properties
    properties.title = f"{company}｜{strategy}策略介绍材料"
    properties.subject = "私募基金投顾推荐报告"
    properties.author = "私募材料智能整理工具"
    properties.last_modified_by = "私募材料智能整理工具"
    properties.keywords = ""
    properties.comments = ""
    properties.category = ""
    properties.created = fixed_time
    properties.modified = fixed_time
    properties.revision = 1


def _normalize_docx_package(data: bytes) -> bytes:
    """固定条目顺序、时间戳和压缩参数，使相同内容生成相同字节。"""
    source = BytesIO(data)
    target = BytesIO()
    with zipfile.ZipFile(source, "r") as archive, zipfile.ZipFile(
        target, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9
    ) as output:
        for name in sorted(archive.namelist()):
            original = archive.getinfo(name)
            info = zipfile.ZipInfo(filename=name, date_time=(1980, 1, 1, 0, 0, 0))
            info.compress_type = zipfile.ZIP_DEFLATED
            info.create_system = 0
            info.external_attr = original.external_attr
            output.writestr(info, archive.read(name))
    return target.getvalue()


def _legacy_research(content: str) -> GeneratedResearch:
    claims = [ResearchClaim(line.strip("# -•"), []) for line in content.splitlines() if line.strip()]
    section = ResearchSection("core_overview", "公司概况", claims[:6])
    return GeneratedResearch("资料未披露。", [], [], [section], [], [])
