"""核心解析、识别、数据库和文档导出测试。"""

from __future__ import annotations

import sqlite3
import zipfile
import hashlib
from io import BytesIO
from pathlib import Path

import pytest
from PIL import Image, ImageDraw
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches

from ai.company_extract import extract_company
from ai.citation_verifier import verify_research_payload
from ai.research_editor import apply_research_edit, editable_facts, promote_blocked_fact
from ai.research_combiner import combine_strategy_research
from ai.module_extract import build_modules
from ai.source_detail_enricher import enrich_research_from_sources
from ai.strategy_extract import extract_strategies
from database.database import init_database, save_modules
from document.txt_generator import compose_natural_brief, generate_txt
from document.performance_chart import build_performance_chart, build_performance_visual
from document.word_generator import generate_word
from file_parser import extract_source_images, parse_file
from file_parser.zip_parser import decode_zip_member_name
from models import SourceChunk, SourceImage
from pages.result_state import get_generated_results, set_active_strategy, store_generated_results


def test_txt_parse_and_identification(tmp_path: Path) -> None:
    """验证文本解析及公司、策略自动识别。"""
    path = tmp_path / "涵德投资路演.txt"
    path.write_text("涵德投资专注量化投资。中证500指数增强采用多因子选股与风险控制。", encoding="utf-8")
    chunks = parse_file(path)
    company = extract_company(chunks)
    assert "涵德投资" in company.name
    assert "中证500指数增强" in extract_strategies(chunks)


def test_legal_company_name_and_confidence_are_preserved() -> None:
    """验证法定全称不被截断，且大量通用投资词不会稀释可信度。"""
    chunks = [
        SourceChunk("*资料来源：北京涵德投资管理有限公司\n公司采用量化投资与组合投资方法。", "涵德策略.pdf", "第 5 页", "PDF"),
        SourceChunk("北京涵德投资管理有限公司投研团队介绍。证券投资、基金投资与风险管理。", "涵德介绍.docx", "段落 2", "Word"),
    ]
    result = extract_company(chunks)
    assert result.name == "北京涵德投资管理有限公司"
    assert result.confidence >= 90


def test_modules_keep_sources() -> None:
    """验证模块整理保留文件名、页码和原文。"""
    chunks = [SourceChunk("策略采用多因子选股，并设置行业偏离约束。", "路演.pptx", "第 15 页", "PPT")]
    modules = build_modules(chunks)
    strategy_module = next(item for item in modules if item.title == "策略逻辑")
    assert strategy_module.sources[0].source_page == "第 15 页"
    assert "多因子" in strategy_module.sources[0].text


def test_zip_path_traversal_is_blocked(tmp_path: Path) -> None:
    """验证 ZIP 路径穿越攻击被拒绝。"""
    path = tmp_path / "unsafe.zip"
    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr("../escape.txt", "unsafe")
    with pytest.raises(ValueError, match="不安全路径"):
        parse_file(path)


def test_gbk_zip_filename_is_repaired() -> None:
    """验证 Windows ZIP 中的 GBK 中文路径可自动恢复。"""
    expected = "【20260618】涵德/【股票版】涵德量化策略全景解析20260603.pdf"
    mojibake = expected.encode("gbk").decode("cp437")
    assert decode_zip_member_name(mojibake) == expected


def test_database_and_exports(tmp_path: Path) -> None:
    """验证 SQLite 保存及 TXT、DOCX 生成。"""
    db_path = tmp_path / "test.db"
    init_database(db_path)
    modules = build_modules([SourceChunk("公司采用风险控制流程。", "介绍.docx", "段落 2", "Word")])
    count = save_modules("测试投资", ["市场中性"], modules, db_path)
    with sqlite3.connect(db_path) as connection:
        assert connection.execute("SELECT COUNT(*) FROM company_info").fetchone()[0] == count
    txt_data, txt_path = generate_txt("测试投资", "市场中性", "公司介绍\n资料未披露", tmp_path)
    docx_data, docx_path = generate_word("测试投资", "市场中性", "## 一、公司基本情况\n资料未披露", tmp_path)
    assert txt_data.startswith(b"\xef\xbb\xbf") and txt_path.exists()
    assert docx_data.startswith(b"PK") and docx_path.exists()


def test_grounding_blocks_unmatched_quotes_and_numbers() -> None:
    """验证逐字引文和数字覆盖两道门槛都能拦截问题事实。"""
    chunks = [
        SourceChunk(
            "北京测试投资管理有限公司成立于2013年，管理规模70亿元，行业偏离不超过3%。",
            "介绍.docx",
            "段落 1",
            "Word",
            "S0001",
        )
    ]
    quote = "北京测试投资管理有限公司成立于2013年，管理规模70亿元，行业偏离不超过3%。"
    payload = {
        "metadata": [],
        "brief_claims": [
            {"text": "公司成立于2013年，管理规模70亿元。", "citations": [{"source_id": "S0001", "quote": quote}]},
            {"text": "行业偏离不超过5%。", "citations": [{"source_id": "S0001", "quote": quote}]},
            {"text": "公司团队稳定。", "citations": [{"source_id": "S0001", "quote": "原文中不存在"}]},
        ],
        "sections": [],
        "performance": [],
    }
    result = verify_research_payload(payload, chunks)
    assert "70亿元" in result.brief
    assert "5%" not in result.brief and "团队稳定" not in result.brief
    blocked_reasons = " ".join(reason for item in result.audit for reason in item.get("reasons", []))
    assert "引文未覆盖数字：5%" in blocked_reasons
    assert "引文不是该位置原文的连续片段" in blocked_reasons


def test_human_fact_edit_is_revalidated_and_rebuilds_research() -> None:
    """验证人工修订可保存，未被原文覆盖的数字会被拒绝，也可删除事实。"""
    source = SourceChunk(
        "北京测试投资成立于2020年，管理规模20亿元。中证500指数增强采用多因子选股。",
        "介绍.docx",
        "段落 1",
        "Word",
        "S0001",
    )
    citation = [{"source_id": "S0001", "quote": source.text}]
    research = verify_research_payload(
        {
            "metadata": [],
            "brief_claims": [{"text": "公司成立于2020年，管理规模20亿元。", "citations": citation}],
            "sections": [
                {
                    "key": "strategy_framework",
                    "claims": [{"text": "中证500指数增强采用多因子选股。", "citations": citation}],
                }
            ],
            "performance": [],
        },
        [source],
    )
    records = editable_facts(research)
    brief = next(item for item in records if item["path"][0] == "brief")
    updated = apply_research_edit(
        research,
        tuple(brief["path"]),
        "公司于2020年成立，披露管理规模为20亿元。",
        [source],
    )
    assert "披露管理规模为20亿元" in updated.brief
    assert any(item.get("manual_edit") for item in updated.audit)

    with pytest.raises(ValueError, match="30"):
        apply_research_edit(
            research,
            tuple(brief["path"]),
            "公司于2020年成立，披露管理规模为30亿元。",
            [source],
        )

    detail = next(item for item in editable_facts(updated) if item["path"][0] == "section")
    deleted = apply_research_edit(updated, tuple(detail["path"]), detail["value"], [source], delete=True)
    assert not next(section for section in deleted.sections if section.key == "strategy_framework").claims


def test_blocked_fact_can_be_corrected_reverified_and_promoted() -> None:
    """验证被拦截事实可改写、重新选原文并进入指定客户章节。"""
    source = SourceChunk(
        "测试投资成立于2020年，管理规模20亿元。",
        "公司介绍.docx",
        "段落 1",
        "Word",
        "S0001",
    )
    research = verify_research_payload(
        {
            "metadata": [],
            "brief_claims": [],
            "sections": [
                {
                    "key": "core_overview",
                    "claims": [
                        {
                            "text": "测试投资成立于2020年，管理规模30亿元。",
                            "citations": [{"source_id": "S0001", "quote": source.text}],
                        }
                    ],
                }
            ],
            "performance": [],
        },
        [source],
    )
    blocked_index = next(index for index, item in enumerate(research.audit) if item.get("status") == "拦截")
    with pytest.raises(ValueError, match="30"):
        promote_blocked_fact(
            research,
            blocked_index,
            "测试投资成立于2020年，管理规模30亿元。",
            ["S0001"],
            [source],
            "section:core_overview",
        )
    promoted = promote_blocked_fact(
        research,
        blocked_index,
        "测试投资成立于2020年，管理规模20亿元。",
        ["S0001"],
        [source],
        "section:core_overview",
    )
    core = next(section for section in promoted.sections if section.key == "core_overview")
    assert core.claims[0].text.endswith("管理规模20亿元。")
    assert promoted.audit[blocked_index]["status"] == "通过"
    assert promoted.audit[blocked_index]["manual_edit"] is True


def test_txt_brief_is_polished_into_natural_multi_strategy_prose() -> None:
    """验证简版先写共同信息，再以一次性小标题分组策略事实。"""
    raw = (
        "共同信息：测试投资成立于2020年。"
        "中证500指数增强：中证500指数增强采用多因子选股。"
        "中证500指数增强：代表产品为500增强1号。"
        "市场中性：市场中性策略采用股指期货对冲。"
    )
    polished = compose_natural_brief(raw, ["中证500指数增强", "市场中性"])
    assert "共同信息：" not in polished
    assert polished.startswith("测试投资成立于2020年。\n\n【中证500指数增强】")
    assert polished.count("【中证500指数增强】") == 1
    assert polished.count("【市场中性】") == 1
    assert "采用多因子选股。代表产品为500增强1号。" in polished
    assert "【市场中性】\n采用股指期货对冲。" in polished
    assert "在中证500指数增强策略方面" not in polished


def test_txt_removes_company_clauses_embedded_in_each_strategy() -> None:
    """验证长句中的公司事实也会去重，同时保留策略独有内容和连续句。"""
    raw = (
        "共同信息：华年投资成立于2023年5月，2024年7月在中基协登记，具备投顾资质。"
        "共同信息：截至2026年6月，公司管理总规模75亿元，其中指数增强20亿元、市场中性10亿元。"
        "共同信息：核心人物薛钰新拥有8年以上从业经验，是市场少数管理过400亿资金的PM之一。"
        "共同信息：公司核心人物为创始人薛钰新，拥有8年以上从业经验，是市场少数管理过400亿资金的PM之一。"
        "中证500指数增强：华年投资成立于2023年5月，2024年7月在中基协登记，具备投顾资质，"
        "截至2026年6月管理总规模75亿元，其中指数增强策略规模20亿元。"
        "该策略采用多因子选股。"
        "市场中性：华年投资成立于2023年5月，具备投顾资质，截至2026年6月管理总规模75亿元，"
        "其中市场中性策略规模10亿元。该策略采用股指期货对冲。"
    )
    polished = compose_natural_brief(raw, ["中证500指数增强", "市场中性"])
    assert polished.count("成立于2023年5月") == 1
    assert polished.count("具备投顾资质") == 1
    assert polished.count("管理总规模75亿元") == 1
    assert polished.count("拥有8年以上从业经验") == 1
    assert polished.count("400亿资金") == 1
    assert "指数增强策略规模20亿元" not in polished
    assert "市场中性策略规模10亿元" not in polished
    assert "【中证500指数增强】\n该策略采用多因子选股。" in polished
    assert "【市场中性】\n该策略采用股指期货对冲。" in polished


def test_multiple_strategy_result_state_keeps_an_active_result() -> None:
    """验证多个策略结果可同时保存并独立切换。"""
    state: dict = {}
    results = {
        "中证500指数增强": {"strategy": "中证500指数增强", "docx_name": "500.docx"},
        "市场中性": {"strategy": "市场中性", "docx_name": "中性.docx"},
    }
    store_generated_results(state, results, "市场中性")
    assert list(get_generated_results(state)) == ["中证500指数增强", "市场中性"]
    assert state["generated"]["strategy"] == "市场中性"
    selected = set_active_strategy(state, "中证500指数增强")
    assert selected["docx_name"] == "500.docx" and state["active_strategy"] == "中证500指数增强"


def test_multiple_strategies_are_combined_into_one_same_format_report(tmp_path: Path) -> None:
    """验证多选策略合并为一个研究对象、一份 TXT 语义和一份同版式 Word。"""
    source = SourceChunk(
        "测试投资成立于2020年。中证500指数增强采用多因子选股，代表产品为500增强1号。市场中性策略采用股指期货对冲，代表产品为中性1号。",
        "策略介绍.docx",
        "段落 1",
        "Word",
        "S0001",
    )
    citation = [{"source_id": "S0001", "quote": source.text}]
    company_claim = "测试投资成立于2020年。"

    def strategy_research(strategy: str, claim: str, product: str):
        return verify_research_payload(
            {
                "metadata": [],
                "brief_claims": [
                    {"text": company_claim, "citations": citation},
                    {"text": claim, "citations": citation},
                ],
                "sections": [
                    {"key": "core_overview", "claims": [{"text": company_claim, "citations": citation}]},
                    {"key": "strategy_framework", "claims": [{"text": claim, "citations": citation}]},
                ],
                "performance": [
                    {"label": "代表产品", "value": product, "citations": citation}
                ],
            },
            [source],
        )

    combined = combine_strategy_research(
        [
            ("中证500指数增强", strategy_research("中证500指数增强", "中证500指数增强采用多因子选股。", "500增强1号")),
            ("市场中性", strategy_research("市场中性", "市场中性策略采用股指期货对冲。", "中性1号")),
        ]
    )
    assert "中证500指数增强" in combined.brief and "市场中性" in combined.brief
    assert [fact.label for fact in combined.performance] == ["代表产品", "代表产品"]
    assert combined.performance[0].value == "中证500指数增强：500增强1号"
    assert combined.performance[1].value == "市场中性：中性1号"
    core = next(section for section in combined.sections if section.key == "core_overview")
    assert len(core.claims) == 1 and core.claims[0].text == "共同信息：测试投资成立于2020年。"
    polished = compose_natural_brief(combined.brief, ["中证500指数增强", "市场中性"])
    assert polished.count("测试投资成立于2020年") == 1
    assert polished.count("【中证500指数增强】") == polished.count("【市场中性】") == 1
    docx_data, docx_path = generate_word(
        "测试投资",
        "中证500指数增强、市场中性",
        combined,
        tmp_path,
        source_chunks=[source],
    )
    assert docx_path.name == "测试投资_中证500指数增强、市场中性_详细介绍.docx"
    document = Document(BytesIO(docx_data))
    document_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "中证500指数增强" in document_text and "市场中性" in document_text
    chapter_titles = [paragraph.text for paragraph in document.paragraphs if "、" in paragraph.text]
    assert sum(title.endswith("中证500指数增强") for title in chapter_titles) == 1
    assert sum(title.endswith("市场中性") for title in chapter_titles) == 1
    assert document_text.count("测试投资成立于2020年") == 1
    assert "中证500指数增强：采用多因子选股" not in document_text
    assert "市场中性：采用股指期货对冲" not in document_text
    assert len(document.tables) == 1
    table_values = [cell.text for row in document.tables[0].rows for cell in row.cells]
    assert "中证500指数增强：500增强1号" not in table_values
    assert "市场中性：中性1号" not in table_values


def test_performance_chart_and_docx_are_deterministic(tmp_path: Path) -> None:
    """验证代表产品图来自原始时间序列，且重复生成的 Word 字节一致。"""
    intro = SourceChunk(
        "代表产品为测试500增强1号。中证500指数增强通过多因子选股获取超额收益。",
        "介绍.docx",
        "段落 1",
        "Word",
        "S0001",
    )
    rows = ["日期 | 测试500增强1号净值 | 中证500基准"]
    rows.extend(f"2025-{month:02d}-01 | {1 + month * 0.02:.2f} | {1 + month * 0.01:.2f}" for month in range(1, 13))
    time_series = SourceChunk("\n".join(rows), "净值.xlsx", "工作表“净值” 行 1-13", "Excel", "S0002")
    payload = {
        "metadata": [],
        "brief_claims": [
            {
                "text": "中证500指数增强通过多因子选股获取超额收益。",
                "citations": [
                    {"source_id": "S0001", "quote": "中证500指数增强通过多因子选股获取超额收益。"}
                ],
            }
        ],
        "sections": [
            {
                "key": "strategy_framework",
                "claims": [
                    {
                        "text": "中证500指数增强通过多因子选股获取超额收益。",
                        "citations": [
                            {"source_id": "S0001", "quote": "中证500指数增强通过多因子选股获取超额收益。"}
                        ],
                    }
                ],
            }
        ],
        "performance": [
            {
                "label": "代表产品",
                "value": "测试500增强1号",
                "citations": [{"source_id": "S0001", "quote": "代表产品为测试500增强1号。"}],
            }
        ],
    }
    research = verify_research_payload(payload, [intro, time_series])
    assert "S0001" not in research.detail_markdown()
    chart = build_performance_chart([intro, time_series], "中证500指数增强", research)
    assert chart is not None and chart["source_id"] == "S0002" and chart["point_count"] == 12
    assert chart["image_data"].startswith(b"\x89PNG")

    first, _ = generate_word(
        "北京测试投资管理有限公司",
        "中证500指数增强",
        research,
        tmp_path,
        source_chunks=[intro, time_series],
        performance_chart=chart,
    )
    second, _ = generate_word(
        "北京测试投资管理有限公司",
        "中证500指数增强",
        research,
        tmp_path,
        source_chunks=[intro, time_series],
        performance_chart=chart,
    )
    assert hashlib.sha256(first).hexdigest() == hashlib.sha256(second).hexdigest()
    with zipfile.ZipFile(BytesIO(first)) as archive:
        package_xml = b"".join(archive.read(name) for name in archive.namelist() if name.endswith(".xml"))
    assert "涵德".encode("utf-8") not in package_xml
    assert b"S0001" not in package_xml and b"S0002" not in package_xml
    assert "资料来源：".encode("utf-8") not in package_xml
    generated_document = Document(BytesIO(first))
    section = generated_document.sections[0]
    assert round(section.page_width.inches, 2) == 8.27
    assert round(section.left_margin.inches, 2) == 1.25
    assert generated_document.paragraphs[0].runs[0].font.name == "宋体"


def test_source_performance_image_is_preferred_and_keeps_origin(tmp_path: Path) -> None:
    """验证原材料中的业绩图优先于重绘图，并保留后台来源定位。"""
    image = Image.new("RGB", (1200, 520), "white")
    draw = ImageDraw.Draw(image)
    draw.line((80, 420, 1120, 80), fill="#1f3864", width=8)
    buffer = BytesIO()
    image.save(buffer, format="PNG")
    source_image = SourceImage(
        buffer.getvalue(),
        "产品路演.pptx",
        "第 16 页",
        "测试500增强1号 净值及超额收益表现 最大回撤 夏普",
        1200,
        520,
        "PPT",
        "I0001",
    )
    research = verify_research_payload(
        {
            "metadata": [],
            "brief_claims": [],
            "sections": [],
            "performance": [
                {
                    "label": "代表产品",
                    "value": "测试500增强1号",
                    "citations": [{"source_id": "S0001", "quote": "代表产品为测试500增强1号。"}],
                }
            ],
        },
        [SourceChunk("代表产品为测试500增强1号。", "介绍.docx", "段落 1", "Word", "S0001")],
    )
    visual = build_performance_visual([source_image], [], "中证500指数增强", research)
    assert visual is not None and visual["visual_type"] == "source_image"
    assert visual["image_id"] == "I0001" and visual["source_file"] == "产品路演.pptx"

    source_docx = tmp_path / "含业绩图.docx"
    source_document = Document()
    source_document.add_paragraph("测试500增强1号净值、超额收益及最大回撤表现")
    source_document.add_picture(BytesIO(buffer.getvalue()), width=Inches(5))
    source_document.save(source_docx)
    images, errors = extract_source_images([source_docx])
    assert not errors and images and images[0].source_file == source_docx.name


def test_sparse_detail_is_enriched_with_fluent_text_and_exact_quotes() -> None:
    """验证详版补充会书面化正文，同时保留逐字原文并避开无关内容。"""
    source_file = "某私募-投顾推荐报告.docx"
    raw = [
        ("公司概况：某私募成立于2020年，有投顾资质，管理规模20亿元。", "段落 3"),
        ("投资团队：总人数15人，核心基金经理有8年从业经验。", "段落 4"),
        ("中证500指增策略（规模10亿元）：", "段落 10"),
        ("选股池：剔除ST股，在全市场约5000只股票中选股。", "段落 11"),
        ("因子层面：量价因子70%、基本面因子30%。", "段落 12"),
        ("换手率：年化双边换手30-40倍，不包含T0。", "段落 13"),
        ("风控方面：使用Barra框架管理行业偏离和风格暴露。", "段落 14"),
        ("代表产品为某500增强1号，超额最大回撤2.8%。", "段落 20"),
        ("市场中性策略使用IC股指期货对冲。", "段落 30"),
        ("基金过往业绩不预示其未来表现。", "段落 40"),
    ]
    chunks = [SourceChunk(text, source_file, page, "Word", f"S{index:04d}") for index, (text, page) in enumerate(raw, 1)]
    research = verify_research_payload({"metadata": [], "brief_claims": [], "sections": [], "performance": []}, chunks)
    enriched = enrich_research_from_sources(research, chunks, "中证500指数增强")
    section_map = {section.key: section for section in enriched.sections}
    selected = [claim for section in enriched.sections for claim in section.claims]
    assert len(section_map["core_overview"].claims) >= 2
    assert len(section_map["strategy_framework"].claims) >= 3
    assert section_map["risk_control"].claims
    assert section_map["performance"].claims
    source_texts = {chunk.source_id: chunk.text for chunk in chunks}
    assert all(claim.citations[0].quote == source_texts[claim.citations[0].source_id] for claim in selected)
    assert any(claim.text != claim.citations[0].quote for claim in selected)
    all_text = " ".join(claim.text for claim in selected)
    assert "在因子研究上" in all_text and "在交易频率上" in all_text
    assert "因子层面：" not in all_text and "换手率：" not in all_text
    assert "市场中性" not in all_text and "过往业绩不预示" not in all_text


def test_word_uses_large_title_sections_small_date_and_performance_summary(tmp_path: Path) -> None:
    """验证大标题、独立章节、小字日期、业绩摘要及客户版隐藏项。"""
    source = SourceChunk(
        "尽调日期为2026年6月3日。公司口径为2026年6月访谈，业绩截止为2026年7月3日。"
        "代表产品为测试500增强1号，年化收益率为23.50%，最大回撤为5.20%。"
        "产品费率以管理人正式合同为准。后续需要确认模型算法和实际胜率。",
        "介绍.docx",
        "段落 1",
        "Word",
        "S0001",
    )
    quote = source.text
    research = verify_research_payload(
        {
            "metadata": [
                {
                    "label": "材料日期",
                    "value": "2026年6月3日",
                    "citations": [{"source_id": "S0001", "quote": quote}],
                },
                {
                    "label": "公司口径",
                    "value": "2026年6月访谈",
                    "citations": [{"source_id": "S0001", "quote": quote}],
                },
                {
                    "label": "业绩截止",
                    "value": "2026年7月3日",
                    "citations": [{"source_id": "S0001", "quote": quote}],
                }
            ],
            "brief_claims": [],
            "sections": [
                {
                    "key": "follow_up",
                    "claims": [
                        {
                            "text": "后续需要确认模型算法和实际胜率。",
                            "citations": [{"source_id": "S0001", "quote": quote}],
                        }
                    ],
                },
                {
                    "key": "other_information",
                    "claims": [
                        {
                            "text": "产品费率以管理人正式合同为准。",
                            "citations": [{"source_id": "S0001", "quote": quote}],
                        }
                    ],
                }
            ],
            "performance": [
                {
                    "label": "代表产品",
                    "value": "测试500增强1号",
                    "citations": [{"source_id": "S0001", "quote": quote}],
                },
                {
                    "label": "年化收益率",
                    "value": "23.50%",
                    "citations": [{"source_id": "S0001", "quote": quote}],
                },
                {
                    "label": "最大回撤",
                    "value": "5.20%",
                    "citations": [{"source_id": "S0001", "quote": quote}],
                },
            ],
        },
        [source],
    )
    docx_data, _ = generate_word("测试投资", "中证500指数增强", research, tmp_path)
    document = Document(BytesIO(docx_data))
    assert document.paragraphs[0].text == "测试投资"
    assert document.paragraphs[0].runs[0].font.size.pt == 18
    assert document.paragraphs[1].text == "中证500指数增强策略介绍材料"
    assert document.paragraphs[1].runs[0].font.size.pt == 20
    metadata = next(paragraph for paragraph in document.paragraphs if paragraph.text.startswith("材料日期："))
    assert all(run.font.size.pt == 9 for run in metadata.runs)
    assert "业绩截止：2026年7月3日" in metadata.text
    assert not any("公司口径" in paragraph.text for paragraph in document.paragraphs)
    assert not any("其他信息" in paragraph.text for paragraph in document.paragraphs)
    assert not any("购买前需进一步确认" in paragraph.text for paragraph in document.paragraphs)
    assert not any("后续需要确认模型算法" in paragraph.text for paragraph in document.paragraphs)
    assert any(paragraph.text == "一、代表产品表现" for paragraph in document.paragraphs)
    assert any("代表产品为测试500增强1号" in paragraph.text for paragraph in document.paragraphs)
    assert any("年化收益率为23.50%" in paragraph.text and "最大回撤为5.20%" in paragraph.text for paragraph in document.paragraphs)

    body = list(document._element.body)
    table_index = next(index for index, element in enumerate(body) if element.tag == qn("w:tbl"))
    narrative_index = next(
        index
        for index, element in enumerate(body)
        if element.tag == qn("w:p")
        and "年化收益率为23.50%" in "".join(node.text or "" for node in element.iter(qn("w:t")))
    )
    assert table_index > narrative_index
